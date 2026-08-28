import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_border(cell, **kwargs):
    """Apply borders to table cells."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'<w:top w:val="{kwargs.get("top", "single")}" w:sz="{kwargs.get("top_sz", "4")}" w:space="0" w:color="{kwargs.get("top_color", "CBD5E1")}"/>\n'
        f'<w:left w:val="{kwargs.get("left", "single")}" w:sz="{kwargs.get("left_sz", "4")}" w:space="0" w:color="{kwargs.get("left_color", "CBD5E1")}"/>\n'
        f'<w:bottom w:val="{kwargs.get("bottom", "single")}" w:sz="{kwargs.get("bottom_sz", "4")}" w:space="0" w:color="{kwargs.get("bottom_color", "CBD5E1")}"/>\n'
        f'<w:right w:val="{kwargs.get("right", "single")}" w:sz="{kwargs.get("right_sz", "4")}" w:space="0" w:color="{kwargs.get("right_color", "CBD5E1")}"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def set_cell_shading(cell, color_hex):
    """Apply background color to table cells."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_page_border(section):
    """Add double border frame to section matching ICTU cover page."""
    sectPr = section._sectPr
    pgBorders = parse_xml(
        f'<w:pgBorders {nsdecls("w")} w:offsetFrom="page">\n'
        f'<w:top w:val="double" w:sz="12" w:space="24" w:color="0F172A"/>\n'
        f'<w:left w:val="double" w:sz="12" w:space="24" w:color="0F172A"/>\n'
        f'<w:bottom w:val="double" w:sz="12" w:space="24" w:color="0F172A"/>\n'
        f'<w:right w:val="double" w:sz="12" w:space="24" w:color="0F172A"/>\n'
        f'</w:pgBorders>'
    )
    sectPr.append(pgBorders)

def build_report():
    doc = Document()

    # --- BASE PAGE SETUP (A4, Standard ICTU Margins) ---
    cover_section = doc.sections[0]
    cover_section.top_margin = Inches(0.79)     # 2.0 cm
    cover_section.bottom_margin = Inches(0.79)  # 2.0 cm
    cover_section.left_margin = Inches(1.18)    # 3.0 cm
    cover_section.right_margin = Inches(0.79)   # 2.0 cm
    add_page_border(cover_section)

    # --- BASE TYPOGRAPHY STYLE ---
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(13)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    normal_style.paragraph_format.line_spacing = 1.3
    normal_style.paragraph_format.space_after = Pt(6)

    # =========================================================================
    # 1. TRANG BÌA (COVER PAGE - EXACT ICTU TEMPLATE PARITY)
    # =========================================================================
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_inst.add_run("TRƯỜNG ĐẠI HỌC CNTT VÀ TRUYỀN THÔNG\n")
    r1.bold = True
    r1.font.size = Pt(13)
    r2 = p_inst.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    r2.bold = True
    r2.underline = True
    r2.font.size = Pt(13)

    p_sp1 = doc.add_paragraph()
    p_sp1.paragraph_format.space_before = Pt(48)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_rep = p_title.add_run("BÁO CÁO DỰ ÁN\n")
    r_rep.bold = True
    r_rep.font.size = Pt(18)
    r_rep.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) # Deep Red

    r_sub = p_title.add_run("HỌC PHẦN ỨNG DỤNG TRÍ TUỆ NHÂN TẠO\n\n")
    r_sub.bold = True
    r_sub.font.size = Pt(15)
    r_sub.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p_topic = doc.add_paragraph()
    p_topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tlabel = p_topic.add_run("Đề tài:\n")
    r_tlabel.font.size = Pt(13)
    r_tlabel.italic = True
    r_tlabel.bold = True
    r_tlabel.underline = True

    r_tname = p_topic.add_run("HỆ THỐNG QUẢN LÝ DỰ ÁN VÀ TIẾN ĐỘ CÔNG VIỆC KOSHI\nCÓ TÍCH HỢP AI\n\n")
    r_tname.bold = True
    r_tname.font.size = Pt(14)

    p_sp2 = doc.add_paragraph()
    p_sp2.paragraph_format.space_before = Pt(18)

    # Authors Roster Table
    t_auth = doc.add_table(rows=6, cols=2)
    t_auth.alignment = WD_TABLE_ALIGNMENT.CENTER
    auth_data = [
        ("Tên nhóm:", "NHÓM 04"),
        ("Nhóm sinh viên thực hiện:", ""),
        ("1. Phạm Minh Tú (#)", ""),
        ("2. Phạm Văn Huynh", ""),
        ("3. Đàm Đức Đôn", ""),
        ("Giảng viên:", "Nguyễn Thị Tuyển")
    ]
    for idx, (col1, col2) in enumerate(auth_data):
        row = t_auth.rows[idx]
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.line_spacing = 1.15
        p0.paragraph_format.space_after = Pt(2)
        r = p0.add_run(col1)
        if idx in [0, 1, 5]:
            r.bold = True
        
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.line_spacing = 1.15
        p1.paragraph_format.space_after = Pt(2)
        r_c2 = p1.add_run(col2)
        if idx == 0:
            r_c2.bold = True
        
        row.cells[0].width = Inches(3.2)
        row.cells[1].width = Inches(2.8)

    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(84)
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("THÁI NGUYÊN, NĂM 2026")
    r_foot.bold = True
    r_foot.font.size = Pt(13)

    # --- BODY SECTION (Separate Section, No Cover Border) ---
    body_section = doc.add_section()
    body_section.top_margin = Inches(0.79)
    body_section.bottom_margin = Inches(0.79)
    body_section.left_margin = Inches(1.18)
    body_section.right_margin = Inches(0.79)

    # =========================================================================
    # 2. MỤC LỤC
    # =========================================================================
    p_toc_head = doc.add_paragraph()
    p_toc_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_toc_head.add_run("MỤC LỤC")
    r.bold = True
    r.font.size = Pt(16)

    toc_items = [
        ("MỤC LỤC", "i"),
        ("PHÂN CÔNG NHIỆM VỤ", "ii"),
        ("MỞ ĐẦU", "iv"),
        ("CHƯƠNG 1. PHÂN TÍCH YÊU CẦU HỆ THỐNG", "1"),
        ("  1.1. Bối cảnh bài toán và lý do phát triển", "1"),
        ("  1.2. Khảo sát hiện trạng và các giải pháp tương tự", "1"),
        ("  1.3. Mô hình Actor và Phân hệ Use Case tổng quát", "2"),
        ("  1.4. Đặc tả yêu cầu chức năng cốt lõi (Functional Requirements)", "3"),
        ("  1.5. Đặc tả yêu cầu phi chức năng (Non-Functional Requirements)", "4"),
        ("  1.6. Xác định bài toán ứng dụng AI và phạm vi tích hợp", "5"),
        ("CHƯƠNG 2. THIẾT KẾ HỆ THỐNG VÀ CƠ SỞ DỮ LIỆU", "6"),
        ("  2.1. Kiến trúc tổng thể hệ thống phân tầng (3-Tier Architecture)", "6"),
        ("  2.2. Thiết kế Cơ sở dữ liệu quan hệ (ERD & Data Dictionary)", "7"),
        ("  2.3. Thiết kế trạng thái công việc và luồng dữ liệu", "8"),
        ("CHƯƠNG 3. TÍCH HỢP AI, THIẾT KẾ PROMPT VÀ ĐÁNH GIÁ", "9"),
        ("  3.1. Kiến trúc luồng gọi AI đa tầng (Cascade Architecture)", "9"),
        ("  3.2. Thiết kế System Prompt và Ma trận tối ưu hóa 3 vòng", "10"),
        ("  3.3. Thuật toán phân tích chuỗi phụ thuộc DAG và đường găng CPM", "11"),
        ("  3.4. Minh chứng ứng dụng AI trong quy trình phát triển", "12"),
        ("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "13"),
        ("TÀI LIỆU THAM KHẢO", "14")
    ]
    for title, pg in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(title)
        if title.startswith("CHƯƠNG") or title in ["MỤC LỤC", "PHÂN CÔNG NHIỆM VỤ", "MỞ ĐẦU", "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "TÀI LIỆU THAM KHẢO"]:
            r_t.bold = True
        r_dots = p.add_run(" " + "." * (68 - len(title)) + " ")
        r_dots.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        r_p = p.add_run(pg)
        r_p.bold = True

    doc.add_page_break()

    # =========================================================================
    # 3. PHÂN CÔNG NHIỆM VỤ (2 TABLES WITH SIGNATURES)
    # =========================================================================
    p_pc_head = doc.add_paragraph()
    p_pc_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_pc_head.add_run("PHÂN CÔNG NHIỆM VỤ")
    r.bold = True
    r.font.size = Pt(16)

    p_t1_title = doc.add_paragraph()
    r = p_t1_title.add_run("Phân công nhiệm vụ theo tiến độ thực hiện:")
    r.bold = True

    t_prog = doc.add_table(rows=6, cols=3)
    t_prog.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_p = ["STT", "Tên nhiệm vụ", "Người thực hiện"]
    for i, h in enumerate(headers_p):
        cell = t_prog.rows[0].cells[i]
        set_cell_shading(cell, "F8FAFC")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True

    prog_rows = [
        ("1", "Phân tích yêu cầu bài toán, lập tài liệu URD, SRS chuẩn ISO/IEC/IEEE 29148", "Phạm Minh Tú"),
        ("2", "Thiết kế kiến trúc hệ thống 3 lớp, thuật toán điều hướng bàn phím 2D và Kahn DAG", "Phạm Minh Tú"),
        ("3", "Thiết kế CSDL quan hệ, bảng project_members, schema DDL và seeder init_db.py", "Phạm Văn Huynh"),
        ("4", "Thiết kế System Prompt AI, ma trận tối ưu 3 vòng và kịch bản kiểm thử API", "Đàm Đức Đôn"),
        ("5", "Tổng hợp báo cáo kỹ thuật KT1, xây dựng DOCX generator và kiểm thử hệ thống", "Cả nhóm")
    ]
    for r_idx, row_data in enumerate(prog_rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = t_prog.rows[r_idx].cells[c_idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val)

    p_space_pc = doc.add_paragraph()
    p_space_pc.paragraph_format.space_before = Pt(12)

    p_t2_title = doc.add_paragraph()
    r = p_t2_title.add_run("Phân công nhiệm vụ theo thành viên thực hiện:")
    r.bold = True

    t_members = doc.add_table(rows=4, cols=4)
    t_members.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_m = ["STT", "Thành Viên", "Nhiệm vụ", "Chữ ký"]
    for i, h in enumerate(headers_m):
        cell = t_members.rows[0].cells[i]
        set_cell_shading(cell, "F8FAFC")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True

    member_rows = [
        ("1", "Phạm Minh Tú\n(#)", "Chịu trách nhiệm kiến trúc tổng thể, mô hình toán học DAG/CPM, URD/SRS ISO standard, điều phối kỹ thuật và bảo mật hệ thống.", ""),
        ("2", "Phạm Văn Huynh", "Phụ trách thiết kế mô hình CSDL, bảng từ điển dữ liệu, kiểm thử ràng buộc SQL và viết kịch bản test backend API.", ""),
        ("3", "Đàm Đức Đôn", "Phụ trách xây dựng và tối ưu ma trận Prompt AI 3 vòng, thiết kế kịch bản test AI, tài liệu báo cáo kỹ thuật.", "")
    ]
    for r_idx, row_data in enumerate(member_rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = t_members.rows[r_idx].cells[c_idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            if c_idx in [0, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_page_break()

    # =========================================================================
    # 4. MỞ ĐẦU
    # =========================================================================
    p_intro_head = doc.add_paragraph()
    p_intro_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_intro_head.add_run("MỞ ĐẦU")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph(
        "Ngày nay, công nghệ thông tin phát triển đồng nghĩa với việc phát triển các phần mềm ứng dụng nhằm tối ưu hóa năng suất lao động và tự động hóa quy trình quản trị. Trong quy trình phát triển phần mềm hiện đại, việc quản lý tiến độ sprint, điều phối công việc và kiểm soát các chuỗi phụ thuộc kỹ thuật đóng vai trò sống còn đối với sự thành bại của dự án."
    )
    doc.add_paragraph(
        "Tuy nhiên, các công cụ quản lý dự án phổ biến hiện nay như Jira, Trello, Asana thường gặp phải các hạn chế lớn về hiệu năng: thời gian phản hồi giao diện chậm, phụ thuộc hoàn toàn vào kết nối mạng Internet, các biểu mẫu nhập liệu đa tầng làm gián đoạn luồng tập trung của lập trình viên, và thiếu sự hỗ trợ tự động hóa thông minh trong việc tổng hợp tiến độ cũng như phân bổ tải công việc."
    )
    doc.add_paragraph(
        "Hệ thống quản lý dự án Koshi (輿) được thiết kế và xây dựng nhằm giải quyết triệt để các hạn chế trên. Với triết lý 'Local-First' và điều hướng tối ưu hóa bàn phím (Vim ergonomics), Koshi mang lại trải nghiệm tương tác với độ trễ dưới 16ms. Đồng thời, việc tích hợp Trí tuệ nhân tạo (AI) đa tầng đóng vai trò như một trợ lý quản lý dự án tự động, hỗ trợ tóm tắt báo cáo tuần, trích xuất biên bản cuộc họp và phân bổ nhân sự cân bằng tải."
    )
    doc.add_paragraph(
        "Báo cáo Dự án Học phần Ứng dụng Trí tuệ Nhân tạo - Bài kiểm tra 1 (KT1) trình bày toàn bộ kết quả phân tích yêu cầu hệ thống, khảo sát hiện trạng, mô hình hóa CSDL và kế hoạch tích hợp AI của Nhóm 04. Nhóm chúng em xin chân thành cảm ơn giảng viên ThS. Nguyễn Thị Tuyển đã tận tình hướng dẫn và định hướng để nhóm hoàn thành tốt đề tài này."
    )

    doc.add_page_break()

    # =========================================================================
    # 5. CHƯƠNG 1: PHÂN TÍCH YÊU CẦU HỆ THỐNG (COMPLETED FULL DEPTH)
    # =========================================================================
    h1 = doc.add_paragraph()
    r = h1.add_run("CHƯƠNG 1. PHÂN TÍCH YÊU CẦU HỆ THỐNG")
    r.bold = True
    r.font.size = Pt(15)

    doc.add_heading("1.1. Bối cảnh bài toán và lý do phát triển", level=2)
    doc.add_paragraph(
        "Koshi được định vị là hệ thống quản lý công việc và tiến độ sprint nội bộ dành cho các đội ngũ kỹ sư phần mềm chuyên sâu. Bối cảnh nghiệp vụ tập trung giải quyết nhu cầu của 3 nhóm đối tượng chính (User Personas):"
    )
    p = doc.add_paragraph()
    p.add_run("• Lead Architect / Senior Engineer: ").bold = True
    p.add_run("Đòi hỏi thao tác 100% bằng bàn phím không dùng chuột (h/j/k/l, Space, i, n, Esc), chuyển đổi tức thì giữa dạng Bảng (Table) và Kanban 2D, phát hiện sớm các điểm nghẽn tiến độ (Critical Path) trên đồ thị phụ thuộc.")
    
    p = doc.add_paragraph()
    p.add_run("• Project Manager (PM) / Tech Lead: ").bold = True
    p.add_run("Cần tự động hóa khâu lập báo cáo tiến độ tuần, tự động trích xuất đầu việc từ biên bản họp văn bản thô và nhận gợi ý phân công công việc dựa trên số điểm độ phức tạp (WIP story points) của từng kỹ sư.")
    
    p = doc.add_paragraph()
    p.add_run("• Field / Mobile Developer: ").bold = True
    p.add_run("Cần khả năng làm việc ngoại tuyến (Offline-first) khi mất kết nối mạng và đồng bộ tự động dữ liệu vào IndexedDB cục bộ với độ trễ phản hồi giao diện dưới 16ms.")

    doc.add_heading("1.2. Khảo sát hiện trạng và các giải pháp tương tự", level=2)
    doc.add_paragraph(
        "Nhóm đã tiến hành khảo sát, so sánh đối chuẩn giữa Koshi và các hệ thống quản lý dự án hàng đầu trên thị trường hiện nay:"
    )

    t_comp = doc.add_table(rows=5, cols=5)
    t_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    comp_headers = ["Tiêu chí đánh giá", "Jira Software", "Trello", "Linear", "Koshi (Đề tài)"]
    for i, h in enumerate(comp_headers):
        cell = t_comp.rows[0].cells[i]
        set_cell_shading(cell, "F1F5F9")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True

    comp_rows = [
        ("Điều hướng phím tắt", "Hạn chế (Phụ thuộc chuột)", "Cơ bản (Kéo thả chuột)", "Khá tốt (Command Menu)", "Toàn diện (Vim 2D grid h/j/k/l)"),
        ("Khả năng Offline", "Không hỗ trợ", "Không hỗ trợ", "Có (Local Cache)", "Toàn diện (Local-first IndexedDB)"),
        ("Phân tích đồ thị DAG", "Phức tạp qua plugin", "Không hỗ trợ", "Không hỗ trợ", "Tích hợp sẵn giải thuật Kahn & CPM"),
        ("Tự động hóa AI PM", "Cơ bản (Tìm kiếm)", "Hạn chế (Power-ups)", "Tự động gán nhãn", "3 luồng AI chuyên sâu + Fallback")
    ]
    for r_idx, row_data in enumerate(comp_rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = t_comp.rows[r_idx].cells[c_idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            if c_idx == 0:
                p.add_run(val).bold = True
            else:
                p.add_run(val)

    doc.add_heading("1.3. Mô hình Actor và Phân hệ Use Case tổng quát", level=2)
    doc.add_paragraph("Hệ thống phân cấp 3 nhóm tác nhân (Actors) với các phân hệ Use Case chính:")
    doc.add_paragraph("1. Guest (Khách vãng lai): Đăng ký tài khoản mới, Đăng nhập hệ thống qua Email/Mật khẩu hoặc Google OAuth2, Xem dữ liệu mẫu thử nghiệm.")
    doc.add_paragraph("2. Team Member (Lập trình viên): Duyệt bảng công việc (Table/Kanban), Chuyển trạng thái vòng tròn (Space), Chỉnh sửa chi tiết công việc (i/Esc), Tạo công việc mới (n), Phân rã mục tiêu bằng AI, Bóc tách Git Diff để tự động đóng ticket.")
    doc.add_paragraph("3. Project Manager (Quản lý dự án): Toàn quyền quản trị thành viên dự án qua `project_members`, Sinh Báo cáo tổng kết tuần (AI Weekly Summary), Bóc tách biên bản họp (AI Meeting Minutes), Cân bằng tải nhân sự (AI Smart Assignment) và trực quan hóa chuỗi phụ thuộc đồ thị (DAG Visualizer).")

    doc.add_heading("1.4. Đặc tả yêu cầu chức năng cốt lõi (Functional Requirements)", level=2)
    fr_list = [
        ("FR-01 [Dual-Mode Views]", "Cho phép chuyển đổi giao diện tức thì giữa Table View mật độ cao và Kanban Board 2D qua phím tắt 'b' với thời gian chuyển cảnh < 16ms."),
        ("FR-02 [Vim Spatial Navigation]", "Hỗ trợ điều hướng bàn phím 2D đầy đủ: j/k duyệt dòng trong Table, h/j/k/l duyệt lưới 4 cột Kanban, Space đổi trạng thái tuần hoàn (TODO -> IN_PROGRESS -> BLOCKED -> DONE -> TODO)."),
        ("FR-03 [Task Detail Inspector]", "Phím Enter mở modal chi tiết, phím 'i' chuyển sang Edit Mode cho phép chỉnh sửa toàn bộ thuộc tính, phím Escape lưu thay đổi và thoát Edit Mode."),
        ("FR-04 [Capture-Phase Escape Trap]", "Bắt sự kiện Escape ở mức window capture phase để đảm bảo đóng modal hoặc hủy chế độ sửa ngay lập tức mà không bị nuốt bởi ô nhập văn bản."),
        ("FR-05 [Project-Scoped RBAC]", "Tách biệt quyền hạn theo từng dự án cụ thể thông qua bảng `project_members`. Cho phép tìm kiếm người dùng trong hệ thống để thêm vào dự án với vai trò OWNER, PM hoặc MEMBER."),
        ("FR-06 [Topological DAG & CPM]", "Sử dụng giải thuật Kahn để phân tích chuỗi phụ thuộc, phát hiện vòng lặp chu trình (Cycle Detection) và tính toán đường găng (Critical Path) cảnh báo điểm nghẽn."),
        ("FR-07 [Autonomous AI PM Workflows]", "Tích hợp 3 chức năng AI chuyên sâu: Tóm tắt tiến độ tuần (3 phần Overview, Blockers, Priorities), Bóc tách biên bản họp ra Action Items, và Gợi ý phân bổ công việc theo kỹ năng và tải trọng WIP."),
        ("FR-08 [Retrospective Work Logging]", "Cho phép tạo công việc mới trực tiếp ở trạng thái DONE để hỗ trợ ghi nhận các tác vụ đột xuất hoặc hotfix mà không cần qua trạng thái trung gian.")
    ]
    for fid, fdesc in fr_list:
        p = doc.add_paragraph()
        r = p.add_run(f"• {fid}: ")
        r.bold = True
        p.add_run(fdesc)

    doc.add_heading("1.5. Đặc tả yêu cầu phi chức năng (Non-Functional Requirements)", level=2)
    nfr_list = [
        ("NFR-01 [Hiệu năng & Độ trễ]", "Mọi thao tác điều hướng bàn phím và render giao diện cục bộ phải phản hồi trong thời gian < 16ms (tương đương 60fps) mà không gây giật khung hình."),
        ("NFR-02 [Độ tin cậy & Concurrency]", "Hệ thống CSDL SQLite phải được cấu hình PRAGMA journal_mode = WAL và busy_timeout = 30000ms nhằm đảm bảo không bị lỗi khóa ghi khi nhiều người dùng thao tác đồng thời."),
        ("NFR-03 [Bảo mật & Xác thực]", "Mật khẩu người dùng được băm an toàn bằng thuật toán Bcrypt. Xác thực API sử dụng JWT HS256 với thời gian hết hạn rõ ràng. Xác thực Google OAuth2 phải kiểm tra chữ ký mật mã nghiêm ngặt."),
        ("NFR-04 [Khả năng phục hồi AI]", "Luồng gọi AI phải áp dụng cơ chế Cascade 3 tầng (Cloud LLM -> Local Ollama -> Heuristic Rule Engine), đảm bảo hệ thống luôn trả về dữ liệu JSON hợp lệ 100% ngay cả khi mất kết nối Internet."),
        ("NFR-05 [Độ tương phản & Công thái học]", "Giao diện tuân thủ bảng màu Slate đơn sắc có độ tương phản cao, chuyển đổi Light/Dark mode với độ trễ 0ms, sử dụng font sans cho tiêu đề và font mono cho mã định danh công việc."),
        ("NFR-06 [Tính toàn vẹn dữ liệu]", "Bật cơ chế PRAGMA foreign_keys = ON trên toàn bộ kết nối cơ sở dữ liệu để đảm bảo các ràng buộc khóa ngoại và xóa xếp tầng (CASCADE) hoạt động chính xác.")
    ]
    for nid, ndesc in nfr_list:
        p = doc.add_paragraph()
        r = p.add_run(f"• {nid}: ")
        r.bold = True
        p.add_run(ndesc)

    doc.add_heading("1.6. Xác định bài toán ứng dụng AI và phạm vi tích hợp", level=2)
    doc.add_paragraph(
        "Trong khuôn khổ học phần Ứng dụng Trí tuệ Nhân tạo, Koshi tập trung ứng dụng các mô hình ngôn ngữ lớn (LLMs) vào 3 bài toán xử lý ngôn ngữ tự nhiên then chốt nhằm nâng cao hiệu suất quản trị dự án:"
    )
    doc.add_paragraph("1. Bài toán Tóm tắt văn bản có cấu trúc (Structured Summarization): Tổng hợp toàn bộ dữ liệu trạng thái công việc trong sprint thành báo cáo súc tích 3 phần (Tổng quan, Điểm nghẽn tiến độ, Ưu tiên tiếp theo).")
    doc.add_paragraph("2. Bài toán Trích xuất thông tin thực thể (Information Extraction): Phân tích văn bản thô từ biên bản cuộc họp để trích xuất danh sách công việc, người chịu trách nhiệm, độ ưu tiên và thời hạn bàn giao.")
    doc.add_paragraph("3. Bài toán Tối ưu hóa phân bổ nguồn lực (Capacity Optimization): Đánh giá ma trận kỹ năng và số điểm độ phức tạp công việc đang thực hiện (WIP points) để đề xuất lập trình viên phù hợp nhất.")

    doc.add_page_break()

    # =========================================================================
    # 6. CHƯƠNG 2: THIẾT KẾ HỆ THỐNG VÀ CƠ SỞ DỮ LIỆU
    # =========================================================================
    h2 = doc.add_paragraph()
    r = h2.add_run("CHƯƠNG 2. THIẾT KẾ HỆ THỐNG VÀ CƠ SỞ DỮ LIỆU")
    r.bold = True
    r.font.size = Pt(15)

    doc.add_heading("2.1. Kiến trúc tổng thể hệ thống phân tầng (3-Tier Architecture)", level=2)
    doc.add_paragraph(
        "Koshi được xây dựng dựa trên mô hình phân tách 3 lớp rõ rệt: Presentation Layer (Vue 3.5 SPA, Tailwind CSS v4, Pinia), Application Core Layer (FastAPI, Python 3.11, Pydantic, SQLAlchemy ORM) và Persistence Layer (SQLite WAL Mode). Hệ thống hỗ trợ Reverse Proxy bảo mật thông qua Caddy và Nginx."
    )

    doc.add_heading("2.2. Thiết kế Cơ sở dữ liệu quan hệ (ERD & Data Dictionary)", level=2)
    doc.add_paragraph("Hệ thống bao gồm 7 bảng thực thể cốt lõi đảm bảo toàn vẹn dữ liệu:")
    
    t_erd = doc.add_table(rows=8, cols=4)
    t_erd.alignment = WD_TABLE_ALIGNMENT.CENTER
    erd_headers = ["Tên bảng", "Khóa chính (PK)", "Khóa ngoại (FK)", "Mô tả nghiệp vụ"]
    for i, h in enumerate(erd_headers):
        cell = t_erd.rows[0].cells[i]
        set_cell_shading(cell, "F1F5F9")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True

    erd_rows = [
        ("users", "id (INTEGER)", "Không", "Lưu trữ tài khoản, mật khẩu băm, Google ID, avatar và kỹ năng."),
        ("projects", "id (INTEGER)", "owner_id -> users(id)", "Lưu trữ thông tin dự án phần mềm và chủ sở hữu dự án."),
        ("project_members", "id (INTEGER)", "project_id, user_id", "Quản lý thành viên và phân quyền trong dự án (OWNER, PM, MEMBER)."),
        ("sprints", "id (INTEGER)", "project_id -> projects(id)", "Quản lý các chu kỳ sprint và trạng thái kích hoạt."),
        ("tasks", "id (INTEGER)", "project_id, sprint_id, assignee_id", "Lưu trữ chi tiết công việc, trạng thái, độ ưu tiên và điểm phức tạp."),
        ("task_dependencies", "id (INTEGER)", "task_id, depends_on_id", "Lưu trữ quan hệ phụ thuộc giữa các công việc phục vụ thuật toán DAG."),
        ("comments", "id (INTEGER)", "task_id, author_id", "Lưu trữ lịch sử trao đổi và bình luận kỹ thuật.")
    ]
    for r_idx, row_data in enumerate(erd_rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = t_erd.rows[r_idx].cells[c_idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            if c_idx in [0, 1]:
                p.add_run(val).bold = True
            else:
                p.add_run(val)

    doc.add_page_break()

    # =========================================================================
    # 7. CHƯƠNG 3: TÍCH HỢP AI, THIẾT KẾ PROMPT VÀ ĐÁNH GIÁ
    # =========================================================================
    h3 = doc.add_paragraph()
    r = h3.add_run("CHƯƠNG 3. TÍCH HỢP AI, THIẾT KẾ PROMPT VÀ ĐÁNH GIÁ")
    r.bold = True
    r.font.size = Pt(15)

    doc.add_heading("3.1. Kiến trúc luồng gọi AI đa tầng (Cascade Architecture)", level=2)
    doc.add_paragraph(
        "Koshi triển khai cơ chế dự phòng 3 tầng: Tầng 1 (Primary Cloud API - Gemini/GPT) -> Tầng 2 (Secondary Local Ollama) -> Tầng 3 (Deterministic Rule Engine). Mô hình đảm bảo tính liên tục của dịch vụ với thời gian timeout được kiểm soát chặt chẽ."
    )

    doc.add_heading("3.2. Ma trận tối ưu hóa Prompt qua 3 vòng thử nghiệm", level=2)
    
    t_pmat = doc.add_table(rows=4, cols=4)
    t_pmat.alignment = WD_TABLE_ALIGNMENT.CENTER
    pm_headers = ["Chức năng AI", "Vòng 1 (Zero-Shot)", "Vòng 2 (Ràng buộc cấu trúc)", "Vòng 3 (Strict JSON & Fallback)"]
    for i, h in enumerate(pm_headers):
        cell = t_pmat.rows[0].cells[i]
        set_cell_shading(cell, "F1F5F9")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True

    pm_rows = [
        ("Weekly Summary", "Văn bản thô dài dòng, sinh đầu việc ảo.", "Chia 3 mục cố định, còn sót markdown thừa.", "Chuẩn hóa Schema Pydantic, nhận diện thẻ Critical Path."),
        ("Meeting Minutes", "Liệt kê gạch đầu dòng, thiếu người nhận việc.", "Bắt buộc trả JSON, đôi khi lỗi cú pháp dài.", "Tích hợp bộ parse JSON tự sửa lỗi + Fallback trích xuất Speaker."),
        ("Smart Workload", "Gợi ý cảm tính, không tính tải trọng.", "Bổ sung điểm Story Point, chưa giải quyết hòa điểm.", "Áp dụng hàm tối ưu hóa chi phí kết hợp chặn trần 8 điểm WIP.")
    ]
    for r_idx, row_data in enumerate(pm_rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = t_pmat.rows[r_idx].cells[c_idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.add_run(val)

    doc.add_page_break()

    # =========================================================================
    # 8. KẾT LUẬN & TÀI LIỆU THAM KHẢO
    # =========================================================================
    p_conc = doc.add_paragraph()
    p_conc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_conc.add_run("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph(
        "Giai đoạn Bài kiểm tra 1 (KT1) đã hoàn thành toàn diện các mục tiêu đề ra: khảo sát thấu đáo hiện trạng các hệ thống quản lý dự án, thiết lập hệ thống yêu cầu chức năng và phi chức năng chuẩn mực, xây dựng mô hình CSDL phân quyền theo dự án và thiết kế cơ chế tích hợp AI đa tầng linh hoạt."
    )
    doc.add_paragraph(
        "Kế hoạch triển khai cho các giai đoạn tiếp theo của đề tài:"
    )
    doc.add_paragraph("• Bài kiểm tra 2 (KT2): Hoàn thiện toàn bộ các API CRUD nghiệp vụ, phân quyền JWT RBAC, tích hợp giao diện phím tắt 2D Kanban và bảng thống kê tiến độ.")
    doc.add_paragraph("• Bài kiểm tra 3 (KT3): Tối ưu hóa sâu luồng gọi AI đa tầng, kiểm thử tự động toàn diện với Pytest và đóng gói triển khai bằng Docker Compose.")

    doc.add_paragraph().paragraph_format.space_before = Pt(24)

    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_ref.add_run("TÀI LIỆU THAM KHẢO")
    r.bold = True
    r.font.size = Pt(16)

    refs = [
        "[1] ISO/IEC/IEEE 29148:2018, 'Systems and software engineering — Life cycle processes — Requirements engineering', IEEE Standards Association, 2018.",
        "[2] Khoa Công nghệ Thông tin - Trường Đại học CNTT & Truyền thông Thái Nguyên (ICTU), 'Đề cương chi tiết và Quy chuẩn Đánh giá Dự án Học phần Ứng dụng Trí tuệ Nhân tạo', Thái Nguyên, 2026.",
        "[3] Tiangolo, S., 'FastAPI Framework Documentation and Architecture Guidelines', Online: [https://fastapi.tiangolo.com/](https://fastapi.tiangolo.com/), 2026.",
        "[4] You, E. et al., 'Vue 3 Composition API & Pinia State Architecture Guide', Online: [https://vuejs.org/](https://vuejs.org/), 2026.",
        "[5] Google Cloud AI & DeepMind, 'Gemini API Technical Guidelines and Prompt Engineering Matrix', Online: [https://ai.google.dev/](https://ai.google.dev/), 2026."
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        p.add_run(ref)

    output_path = os.path.expanduser("~/koshi/nhom4.docx")
    doc.save(output_path)
    print(f"[✓] Document compiled successfully to: {output_path}")

if __name__ == "__main__":
    build_report()
