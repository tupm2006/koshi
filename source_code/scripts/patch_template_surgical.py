import os
import sys
from docx import Document

SRC_TEMPLATE = os.path.expanduser("~/Documents/BAI DU AN_UNG DUNG AI.docx")
DEST_PATH = os.path.expanduser("~/koshi/nhom4.docx")

if not os.path.exists(SRC_TEMPLATE):
    print(f"[!] Error: Template file not found at {SRC_TEMPLATE}", file=sys.stderr)
    sys.exit(1)

doc = Document(SRC_TEMPLATE)

def replace_in_run_list(paragraphs, search_text, replace_text):
    """Replaces text cleanly across runs inside paragraph collections."""
    for p in paragraphs:
        if search_text in p.text:
            # Check single run match first to preserve exact font/color styling
            for r in p.runs:
                if search_text in r.text:
                    r.text = r.text.replace(search_text, replace_text)
                    return
            # Fallback for split runs: rewrite paragraph text
            p.text = p.text.replace(search_text, replace_text)

# -----------------------------------------------------------------------------
# 1. SURGICAL COVER PAGE & METADATA REPLACEMENTS
# -----------------------------------------------------------------------------
METADATA_MAP = [
    ("HỆ THỐNG QUẢN LÝ BÁN HÀNG CÓ TÍCH HỢP AI", "HỆ THỐNG QUẢN LÝ DỰ ÁN VÀ TIẾN ĐỘ CÔNG VIỆC KOSHI CÓ TÍCH HỢP AI"),
    ("Hệ thống quản lý bán hàng có tích hợp AI", "Hệ thống quản lý dự án và tiến độ công việc Koshi có tích hợp AI"),
    ("NHÓM 01", "NHÓM 04"),
    ("Nhóm 01", "Nhóm 04"),
    ("Lương Văn Trà (#)", "Phạm Minh Tú (#)"),
    ("Lương Văn Trà", "Phạm Minh Tú"),
    ("Tráng A Thư", "Phạm Văn Huynh"),
    ("Nguyễn Mạnh Dương", "Đàm Đức Đôn")
]

for p in doc.paragraphs:
    for old_s, new_s in METADATA_MAP:
        replace_in_run_list([p], old_s, new_s)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for old_s, new_s in METADATA_MAP:
                replace_in_run_list(cell.paragraphs, old_s, new_s)

# -----------------------------------------------------------------------------
# 2. POPULATE THE TWO TASK ALLOCATION TABLES
# -----------------------------------------------------------------------------
for table in doc.tables:
    t_text = " ".join(c.text for row in table.rows for c in row.cells)
    
    # Bảng 1: Phân công theo tiến độ
    if "Tên nhiệm vụ" in t_text and "Người thực hiện" in t_text:
        tasks = [
            ("1", "Phân tích yêu cầu bài toán, khảo sát hiện trạng, lập URD & SRS", "Phạm Minh Tú"),
            ("2", "Thiết kế mô hình Actor, Use Case và các yêu cầu chức năng/phi chức năng", "Phạm Minh Tú"),
            ("3", "Xác định mô hình CSDL quan hệ sơ bộ và các thực thể cốt lõi", "Phạm Văn Huynh"),
            ("4", "Xác định phạm vi ứng dụng AI và khảo sát các luồng xử lý NLP", "Đàm Đức Đôn"),
            ("5", "Tổng hợp báo cáo kỹ thuật Chương 1 (KT1) và rà soát tài liệu", "Cả nhóm")
        ]
        for idx, (stt, tname, assignee) in enumerate(tasks, start=1):
            if idx < len(table.rows):
                row = table.rows[idx]
                if len(row.cells) >= 3:
                    row.cells[0].paragraphs[0].text = stt
                    row.cells[1].paragraphs[0].text = tname
                    row.cells[2].paragraphs[0].text = assignee
        # Clear extra rows in template table if any
        for idx in range(len(tasks) + 1, len(table.rows)):
            for cell in table.rows[idx].cells:
                cell.text = ""

    # Bảng 2: Phân công theo thành viên & chữ ký
    elif "Thành Viên" in t_text and "Chữ ký" in t_text:
        members = [
            ("1", "Phạm Minh Tú (#)", "Trưởng nhóm: Phân tích yêu cầu hệ thống, mô hình Use Case, URD/SRS ISO standard và điều phối kỹ thuật.", ""),
            ("2", "Phạm Văn Huynh", "Phụ trách khảo sát hiện trạng, phân tích yêu cầu phi chức năng và thiết kế sơ bộ thực thể CSDL.", ""),
            ("3", "Đàm Đức Đôn", "Phụ trách xác định phạm vi tích hợp AI, khảo sát mô hình LLM và tổng hợp báo cáo kỹ thuật.", "")
        ]
        for idx, (stt, name, desc, sig) in enumerate(members, start=1):
            if idx < len(table.rows):
                row = table.rows[idx]
                if len(row.cells) >= 4:
                    row.cells[0].paragraphs[0].text = stt
                    row.cells[1].paragraphs[0].text = name
                    row.cells[2].paragraphs[0].text = desc
                    row.cells[3].paragraphs[0].text = sig

# -----------------------------------------------------------------------------
# 3. STRIP CD DISC NOTE
# -----------------------------------------------------------------------------
cd_markers = [
    "Ghi chú:",
    "Mỗi nhóm in 1 quyển báo cáo, kèm theo 1 đĩa CD",
    "Nội dung đĩa CD bao gồm",
    "Hướng dẫn cái đặt chương trình thử nghiệm",
    "Nhãn đĩa CD theo mẫu sau",
    "CNTT K23D"
]
for p in doc.paragraphs:
    if any(m.lower() in p.text.lower() for m in cd_markers):
        p.text = ""
        for elem in p._element.xpath('.//*[local-name()="drawing" or local-name()="shape" or local-name()="group"]'):
            if elem.getparent() is not None:
                elem.getparent().remove(elem)

# -----------------------------------------------------------------------------
# 4. POPULATE ONLY CHAPTER 1 (STRICT SCOPE LOCK)
# -----------------------------------------------------------------------------
ch1_body = """1.1. Bối cảnh bài toán và lý do phát triển
Koshi được định vị là hệ thống quản lý công việc và tiến độ sprint nội bộ dành cho các đội ngũ kỹ sư phần mềm chuyên sâu. Bối cảnh nghiệp vụ tập trung giải quyết nhu cầu của 3 nhóm đối tượng chính (User Personas):
• Lead Architect / Senior Engineer: Đòi hỏi thao tác 100% bằng bàn phím không dùng chuột (h/j/k/l, Space, i, n, Esc), chuyển đổi tức thì giữa dạng Bảng (Table) và Kanban 2D, phát hiện sớm các điểm nghẽn tiến độ (Critical Path) trên đồ thị phụ thuộc.
• Project Manager (PM) / Tech Lead: Cần tự động hóa khâu lập báo cáo tiến độ tuần, tự động trích xuất đầu việc từ biên bản họp văn bản thô và nhận gợi ý phân công công việc dựa trên số điểm độ phức tạp (WIP story points) của từng kỹ sư.
• Field / Mobile Developer: Cần khả năng làm việc ngoại tuyến (Offline-first) khi mất kết nối mạng và đồng bộ tự động dữ liệu vào IndexedDB cục bộ với độ trễ phản hồi giao diện dưới 16ms.

1.2. Khảo sát hiện trạng và các giải pháp tương tự
Khảo sát đối chuẩn giữa Koshi và các hệ thống quản lý dự án hàng đầu hiện nay:
• Jira Software: Hệ thống hoàn chỉnh nhưng cồng kềnh, thời gian phản hồi giao diện chậm, phụ thuộc 100% vào kết nối mạng, thao tác bàn phím hạn chế.
• Trello: Trực quan, đơn giản nhưng thiếu khả năng phân tích đồ thị phụ thuộc DAG, không tối ưu cho luồng làm việc kỹ thuật cao.
• Linear: Giao diện hiện đại, hỗ trợ phím tắt tốt nhưng chưa hỗ trợ phân rã mục tiêu chuyên sâu bằng AI và không có cơ chế phân tích đường găng CPM tự động.
• Koshi (Hệ thống đề xuất): Kết hợp triết lý Local-First, điều hướng bàn phím 2D (Vim ergonomics), phân tích chuỗi phụ thuộc bằng giải thuật Kahn kết hợp 3 luồng trợ lý AI quản trị dự án có cơ chế dự phòng 3 tầng.

1.3. Mô hình Actor và Phân hệ Use Case tổng quát
Hệ thống phân cấp 3 nhóm tác nhân (Actors) với các phân hệ Use Case chính:
1. Guest (Khách vãng lai): Đăng ký tài khoản mới, Đăng nhập hệ thống qua Email/Mật khẩu hoặc Google OAuth2, Xem dữ liệu mẫu thử nghiệm.
2. Team Member (Lập trình viên): Duyệt bảng công việc (Table/Kanban), Chuyển trạng thái vòng tròn (Space), Chỉnh sửa chi tiết công việc (i/Esc), Tạo công việc mới (n), Phân rã mục tiêu bằng AI, Bóc tách Git Diff để tự động đóng ticket.
3. Project Manager (Quản lý dự án): Quản trị thành viên dự án qua project_members, Sinh Báo cáo tổng kết tuần (AI Weekly Summary), Bóc tách biên bản họp (AI Meeting Minutes), Cân bằng tải nhân sự (AI Smart Assignment) và trực quan hóa chuỗi phụ thuộc đồ thị (DAG Visualizer).

1.4. Đặc tả yêu cầu chức năng cốt lõi (Functional Requirements)
• FR-01 [Dual-Mode Views]: Chuyển đổi giao diện tức thì giữa Table View mật độ cao và Kanban Board 2D qua phím tắt 'b' với thời gian < 16ms.
• FR-02 [Vim Spatial Navigation]: Hỗ trợ điều hướng bàn phím đầy đủ: j/k duyệt dòng trong Table, h/j/k/l duyệt lưới Kanban, Space đổi trạng thái tuần hoàn (TODO -> IN_PROGRESS -> BLOCKED -> DONE -> TODO).
• FR-03 [Task Detail Inspector]: Phím Enter mở modal chi tiết, phím 'i' chuyển sang Edit Mode cho phép chỉnh sửa toàn bộ thuộc tính, phím Escape lưu thay đổi và thoát Edit Mode.
• FR-04 [Capture-Phase Escape Trap]: Bắt sự kiện Escape ở mức window capture phase để đảm bảo đóng modal hoặc hủy chế độ sửa ngay lập tức mà không bị nuốt bởi ô nhập văn bản.
• FR-05 [Project-Scoped RBAC]: Tách biệt quyền hạn theo từng dự án cụ thể thông qua bảng project_members. Cho phép tìm kiếm người dùng trong hệ thống để thêm vào dự án với vai trò OWNER, PM hoặc MEMBER.
• FR-06 [Topological DAG & CPM]: Sử dụng giải thuật Kahn để phân tích chuỗi phụ thuộc, phát hiện vòng lặp chu trình (Cycle Detection) và tính toán đường găng (Critical Path) cảnh báo điểm nghẽn.
• FR-07 [Autonomous AI PM Workflows]: Tích hợp 3 chức năng AI chuyên sâu: Tóm tắt tiến độ tuần (3 phần Overview, Blockers, Priorities), Bóc tách biên bản họp ra Action Items, và Gợi ý phân bổ công việc theo kỹ năng và tải trọng WIP.
• FR-08 [Retrospective Work Logging]: Cho phép tạo công việc mới trực tiếp ở trạng thái DONE để hỗ trợ ghi nhận các tác vụ đột xuất hoặc hotfix mà không cần qua trạng thái trung gian.

1.5. Đặc tả yêu cầu phi chức năng (Non-Functional Requirements)
• NFR-01 [Hiệu năng & Độ trễ]: Mọi thao tác điều hướng bàn phím và render giao diện cục bộ phải phản hồi trong thời gian < 16ms (tương đương 60fps).
• NFR-02 [Độ tin cậy & Concurrency]: Hệ thống CSDL SQLite phải được cấu hình PRAGMA journal_mode = WAL và busy_timeout = 30000ms nhằm đảm bảo không bị lỗi khóa ghi khi nhiều người dùng thao tác đồng thời.
• NFR-03 [Bảo mật & Xác thực]: Mật khẩu người dùng được băm an toàn bằng Bcrypt. Xác thực API sử dụng JWT HS256. Xác thực Google OAuth2 kiểm tra chữ ký mật mã nghiêm ngặt.
• NFR-04 [Khả năng phục hồi AI]: Luồng gọi AI áp dụng cơ chế Cascade 3 tầng (Cloud LLM -> Local Ollama -> Heuristic Rule Engine), đảm bảo hệ thống luôn trả về dữ liệu JSON hợp lệ ngay cả khi mất kết nối Internet.
• NFR-05 [Độ tương phản & Công thái học]: Giao diện tuân thủ bảng màu Slate đơn sắc có độ tương phản cao, chuyển đổi Light/Dark mode với độ trễ 0ms.
• NFR-06 [Tính toàn vẹn dữ liệu]: Bật cơ chế PRAGMA foreign_keys = ON trên toàn bộ kết nối cơ sở dữ liệu để đảm bảo các ràng buộc khóa ngoại và xóa xếp tầng (CASCADE) hoạt động chính xác.

1.6. Xác định bài toán ứng dụng AI và phạm vi tích hợp
Trong khuôn khổ đề tài, Koshi tập trung ứng dụng các mô hình ngôn ngữ lớn (LLMs) vào 3 bài toán xử lý ngôn ngữ tự nhiên then chốt:
1. Bài toán Tóm tắt văn bản có cấu trúc (Structured Summarization): Tổng hợp dữ liệu trạng thái công việc trong sprint thành báo cáo súc tích 3 phần (Tổng quan, Điểm nghẽn tiến độ, Ưu tiên tiếp theo).
2. Bài toán Trích xuất thông tin thực thể (Information Extraction): Phân tích văn bản thô từ biên bản cuộc họp để trích xuất danh sách công việc, người chịu trách nhiệm, độ ưu tiên và thời hạn bàn giao.
3. Bài toán Tối ưu hóa phân bổ nguồn lực (Capacity Optimization): Đánh giá ma trận kỹ năng và số điểm độ phức tạp công việc đang thực hiện (WIP points) để đề xuất lập trình viên phù hợp nhất."""

# Find "CHƯƠNG 1" in template (body section, after front matter) and replace only its placeholder
for idx, p in enumerate(doc.paragraphs):
    if "CHƯƠNG 1" in p.text.upper() and idx > 15:
        for next_idx in range(idx + 1, min(idx + 5, len(doc.paragraphs))):
            next_p = doc.paragraphs[next_idx]
            if "1.1." in next_p.text or "Thuật ngữ" in next_p.text or "Nội dung" in next_p.text:
                next_p.text = ch1_body
                break
        break

# DO NOT TOUCH CHAPTER 2, CHAPTER 3, OR CONCLUSION (Leave as template placeholders)

# -----------------------------------------------------------------------------
# 5. SAVE MODIFIED TEMPLATE
# -----------------------------------------------------------------------------
doc.save(DEST_PATH)
print(f"[✓] Successfully mutated template in-place and saved to: {DEST_PATH}")

if __name__ == "__main__":
    pass
