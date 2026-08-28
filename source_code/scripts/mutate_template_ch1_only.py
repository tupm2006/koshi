import os
from docx import Document

TEMPLATE_PATH = os.path.expanduser("~/Documents/BAI DU AN_UNG DUNG AI.docx")
OUTPUT_PATH = os.path.expanduser("~/koshi/nhom4.docx")

if not os.path.exists(TEMPLATE_PATH):
    raise FileNotFoundError(f"Template not found at {TEMPLATE_PATH}. Ensure the file exists.")

doc = Document(TEMPLATE_PATH)

def replace_text_in_paragraph(p, search_text, replace_text):
    if search_text in p.text:
        p.text = p.text.replace(search_text, replace_text)

def replace_text_in_table(table, search_text, replace_text):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if search_text in p.text:
                    p.text = p.text.replace(search_text, replace_text)

print("[1/6] Updating Cover Page metadata in Table 0...")
# Table 0 is the Cover Page Table in the template
t0 = doc.tables[0]
replace_text_in_table(t0, "HỆ THỐNG QUẢN LÝ BÁN HÀNG CÓ TÍCH HỢP AI", "HỆ THỐNG QUẢN LÝ DỰ ÁN VÀ TIẾN ĐỘ CÔNG VIỆC KOSHI CÓ TÍCH HỢP AI")
replace_text_in_table(t0, "NHÓM 01", "NHÓM 04")
replace_text_in_table(t0, "Lương Văn Trà (#)", "Phạm Minh Tú (#)")
replace_text_in_table(t0, "Lương Văn Trà", "Phạm Minh Tú")
replace_text_in_table(t0, "Tráng A Thư", "Phạm Văn Huynh")
replace_text_in_table(t0, "Nguyễn Mạnh Dương", "Đàm Đức Đôn")

print("[2/6] Updating Table of Contents...")
doc.paragraphs[6].text = "CHƯƠNG 1. PHÂN TÍCH YÊU CẦU HỆ THỐNG\t1"
doc.paragraphs[7].text = (
    "  1.1. Bối cảnh bài toán và lý do phát triển\t1\n"
    "  1.2. Khảo sát hiện trạng và các giải pháp tương tự\t1\n"
    "  1.3. Mô hình Actor và Phân hệ Use Case tổng quát\t2\n"
    "  1.4. Đặc tả yêu cầu chức năng cốt lõi (FR)\t3\n"
    "  1.5. Đặc tả yêu cầu phi chức năng (NFR)\t4\n"
    "  1.6. Xác định bài toán ứng dụng AI và phạm vi tích hợp\t5"
)
doc.paragraphs[8].text = "CHƯƠNG 2. THIẾT KẾ HỆ THỐNG (BÀI KIỂM TRA 2)\t6"
doc.paragraphs[9].text = "  2.1. Nội dung thiết kế kiến trúc và CSDL (KT2)\t6"
doc.paragraphs[10].text = "CHƯƠNG 3. TÍCH HỢP AI VÀ ĐÁNH GIÁ (BÀI KIỂM TRA 3)\t7"
doc.paragraphs[11].text = "  3.1. Thử nghiệm AI và đánh giá (KT3)\t7"
doc.paragraphs[12].text = "KẾT LUẬN\t8"
doc.paragraphs[13].text = "TÀI LIỆU THAM KHẢO\t9"

print("[3/6] Updating Task Allocation Tables (KT1 Scope)...")
# Table 1: Progress allocation (doc.tables[1])
t1 = doc.tables[1]
progress_data = [
    ("1", "Phân tích yêu cầu bài toán, khảo sát hiện trạng, lập URD & SRS chuẩn ISO/IEC/IEEE 29148", "Phạm Minh Tú"),
    ("2", "Thiết kế mô hình Actor, Use Case và các đặc tả yêu cầu chức năng/phi chức năng", "Phạm Minh Tú"),
    ("3", "Xác định mô hình CSDL quan hệ sơ bộ, phân quyền dự án và các thực thể cốt lõi", "Phạm Văn Huynh"),
    ("4", "Xác định phạm vi ứng dụng AI và khảo sát các luồng xử lý NLP/LLM", "Đàm Đức Đôn"),
    ("5", "Tổng hợp báo cáo kỹ thuật Chương 1 (KT1), kiểm thử và rà soát tài liệu", "Cả nhóm")
]
for r_idx, (stt, task_name, assignee) in enumerate(progress_data, start=1):
    if r_idx < len(t1.rows):
        row = t1.rows[r_idx]
        if len(row.cells) >= 3:
            row.cells[0].paragraphs[0].text = stt
            row.cells[1].paragraphs[0].text = task_name
            row.cells[2].paragraphs[0].text = assignee
for r_idx in range(len(progress_data) + 1, len(t1.rows)):
    for cell in t1.rows[r_idx].cells:
        cell.text = ""

# Table 2: Member tasks & signatures (doc.tables[2])
t2 = doc.tables[2]
member_data = [
    ("1", "Phạm Minh Tú\n(#)", "Trưởng nhóm: Phân tích yêu cầu hệ thống, mô hình Use Case, URD/SRS ISO standard và điều phối kỹ thuật.", ""),
    ("2", "Phạm Văn Huynh", "Phụ trách khảo sát hiện trạng, phân tích yêu cầu phi chức năng và thiết kế sơ bộ thực thể CSDL.", ""),
    ("3", "Đàm Đức Đôn", "Phụ trách xác định phạm vi tích hợp AI, khảo sát mô hình LLM và tổng hợp báo cáo kỹ thuật.", "")
]
for r_idx, (stt, name, role_desc, sig) in enumerate(member_data, start=1):
    if r_idx < len(t2.rows):
        row = t2.rows[r_idx]
        if len(row.cells) >= 4:
            row.cells[0].paragraphs[0].text = stt
            row.cells[1].paragraphs[0].text = name
            row.cells[2].paragraphs[0].text = role_desc
            row.cells[3].paragraphs[0].text = sig

print("[4/6] Updating Introduction and Stripping CD Disc page...")
doc.paragraphs[24].text = "MỞ ĐẦU"
doc.paragraphs[25].text = (
    "Ngày nay, công nghệ thông tin phát triển đồng nghĩa với việc phát triển các phần mềm ứng dụng "
    "nhằm tối ưu hóa năng suất lao động và tự động hóa quy trình quản trị. Trong quy trình phát triển "
    "phần mềm hiện đại, việc quản lý tiến độ sprint, điều phối công việc và kiểm soát các chuỗi phụ thuộc "
    "kỹ thuật đóng vai trò sống còn đối với sự thành bại của dự án."
)
doc.paragraphs[26].text = (
    "Hệ thống quản lý dự án Koshi (輿) được thiết kế và xây dựng nhằm giải quyết các bài toán về hiệu năng "
    "và tự động hóa quản lý. Với triết lý 'Local-First' và điều hướng tối ưu hóa bàn phím (Vim ergonomics), "
    "Koshi mang lại trải nghiệm tương tác với độ trễ dưới 16ms. Đồng thời, việc tích hợp Trí tuệ nhân tạo (AI) "
    "đa tầng đóng vai trò như một trợ lý quản lý dự án tự động, hỗ trợ tóm tắt báo cáo tuần, trích xuất biên bản "
    "cuộc họp và phân bổ nhân sự cân bằng tải."
)
doc.paragraphs[27].text = "Nhóm 04 chúng em xin chân thành cảm ơn giảng viên ThS. Nguyễn Thị Tuyển đã tận tình hướng dẫn để nhóm hoàn thành tốt Bài kiểm tra 1 (KT1) này!"

# Clear CD Disc guidelines and placeholders (P28 to P37)
for i in range(28, 38):
    doc.paragraphs[i].text = ""

print("[5/6] Populating ONLY Chapter 1 (Strict Scope Lock)...")
ch1_full_content = """1.1. Bối cảnh bài toán và lý do phát triển
Koshi được định vị là hệ thống quản lý công việc và tiến độ sprint nội bộ dành cho các đội ngũ kỹ sư phần mềm chuyên sâu. Bối cảnh nghiệp vụ tập trung giải quyết nhu cầu của 3 nhóm đối tượng chính (User Personas):
• Lead Architect / Senior Engineer: Đòi hỏi thao tác 100% bằng bàn phím không dùng chuột (h/j/k/l, Space, i, n, Esc), chuyển đổi tức thì giữa dạng Bảng (Table) và Kanban 2D, phát hiện sớm các điểm nghẽn tiến độ (Critical Path) trên đồ thị phụ thuộc.
• Project Manager (PM) / Tech Lead: Cần tự động hóa khâu lập báo cáo tiến độ tuần, tự động trích xuất đầu việc từ biên bản họp văn bản thô và nhận gợi ý phân công công việc dựa trên số điểm độ phức tạp (WIP story points) của từng kỹ sư.
• Field / Mobile Developer: Cần khả năng làm việc ngoại tuyến (Offline-first) khi mất kết nối mạng và đồng bộ tự động dữ liệu vào IndexedDB cục bộ với độ trễ phản hồi giao diện dưới 16ms.

1.2. Khảo sát hiện trạng và các giải pháp tương tự
Khảo sát đối chuẩn giữa Koshi và các hệ thống quản lý dự án hàng đầu hiện nay:
• Jira Software: Hệ thống hoàn chỉnh nhưng cồng kềnh, thời gian phản hồi giao diện chậm, phụ thuộc 100% vào kết nối mạng, thao tác bàn phím hạn chế.
• Trello: Trực quan, đơn giản nhưng thiếu khả năng phân tích đồ thị phụ thuộc DAG, không tối ưu cho luồng làm việc kỹ thuật cao.
• Linear: Giao diện hiện đại, hỗ trợ phím tắt tốt nhưng chưa hỗ trợ phân rã mục tiêu chuyên sâu bằng AI và không có cơ chế phân tích đường găng CPM tự động.
• Koshi (Hệ thống đề xuất): Kết hợp triết lý Local-First, điều hướng bàn phím 2D (Vim ergonomics), phân tích chuỗi phụ thuộc bằng giải thuật Kahn kết hợp 3 luồng trợ lý AI quản trị dự án có cơ chế dự phòng 3 tầng.

1.3. Mô hình Actor và Phân hệ Use Case tổng quát
Hệ thống phân cấp 3 nhóm tác nhân (Actors) với các phân hệ Use Case chính:
1. Guest (Khách vãng lai): Đăng ký tài khoản mới, Đăng nhập hệ thống qua Email/Mật khẩu hoặc Google OAuth2, Xem dữ liệu mẫu thử nghiệm.
2. Team Member (Lập trình viên): Duyệt bảng công việc (Table/Kanban), Chuyển trạng thái vòng tròn (Space), Chỉnh sửa chi tiết công việc (i/Esc), Tạo công việc mới (n), Phân rã mục tiêu bằng AI, Bóc tách Git Diff để tự động đóng ticket.
3. Project Manager (Quản lý dự án): Quản trị thành viên dự án qua project_members, Sinh Báo cáo tổng kết tuần (AI Weekly Summary), Bóc tách biên bản họp (AI Meeting Minutes), Cân bằng tải nhân sự (AI Smart Assignment) và trực quan hóa chuỗi phụ thuộc đồ thị (DAG Visualizer).

1.4. Đặc tả yêu cầu chức năng cốt lõi (Functional Requirements)
• FR-01 [Dual-Mode Views]: Chuyển đổi giao diện tức thì giữa Table View mật độ cao và Kanban Board 2D qua phím tắt 'b' với thời gian < 16ms.
• FR-02 [Vim Spatial Navigation]: Hỗ trợ điều hướng bàn phím đầy đủ: j/k duyệt dòng trong Table, h/j/k/l duyệt lưới Kanban, Space đổi trạng thái tuần hoàn (TODO -> IN_PROGRESS -> BLOCKED -> DONE -> TODO).
• FR-03 [Task Detail Inspector]: Phím Enter mở modal chi tiết, phím 'i' chuyển sang Edit Mode cho phép chỉnh sửa toàn bộ thuộc tính, phím Escape lưu thay đổi và thoát Edit Mode.
• FR-04 [Capture-Phase Escape Trap]: Bắt sự kiện Escape ở mức window capture phase để đảm bảo đóng modal hoặc hủy chế độ sửa ngay lập tức mà không bị nuốt bởi ô nhập văn bản.
• FR-05 [Project-Scoped RBAC]: Tách biệt quyền hạn theo từng dự án cụ thể thông qua bảng project_members. Cho phép tìm kiếm người dùng trong hệ thống để thêm vào dự án với vai trò OWNER, PM hoặc MEMBER.
• FR-06 [Topological DAG & CPM]: Sử dụng giải thuật Kahn để phân tích chuỗi phụ thuộc, phát hiện vòng lặp chu trình (Cycle Detection) và tính toán đường găng (Critical Path) cảnh báo điểm nghẽn.
• FR-07 [Autonomous AI PM Workflows]: Tích hợp 3 chức năng AI chuyên sâu: Tóm tắt tiến độ tuần (3 phần Overview, Blockers, Priorities), Bóc tách biên bản họp ra Action Items, và Gợi ý phân bổ công việc theo kỹ năng và tải trọng WIP.
• FR-08 [Retrospective Work Logging]: Cho phép tạo công việc mới trực tiếp ở trạng thái DONE để hỗ trợ ghi nhận các tác vụ đột xuất hoặc hotfix mà không cần qua trạng thái trung gian.

1.5. Đặc tả yêu cầu phi chức năng (Non-Functional Requirements)
• NFR-01 [Hiệu năng & Độ trễ]: Mọi thao tác điều hướng bàn phím và render giao diện cục bộ phải phản hồi trong thời gian < 16ms (tương đương 60fps).
• NFR-02 [Độ tin cậy & Concurrency]: Hệ thống CSDL SQLite phải được cấu hình PRAGMA journal_mode = WAL và busy_timeout = 30000ms nhằm đảm bảo không bị lỗi khóa ghi khi nhiều người dùng thao tác đồng thời.
• NFR-03 [Bảo mật & Xác thực]: Mật khẩu người dùng được băm an toàn bằng Bcrypt. Xác thực API sử dụng JWT HS256. Xác thực Google OAuth2 kiểm tra chữ ký mật mã nghiêm ngặt.
• NFR-04 [Khả năng phục hồi AI]: Luồng gọi AI áp dụng cơ chế Cascade 3 tầng (Cloud LLM -> Local Ollama -> Heuristic Rule Engine), đảm bảo hệ thống luôn trả về dữ liệu JSON hợp lệ ngay cả khi mất kết nối Internet.
• NFR-05 [Độ tương phản & Công thái học]: Giao diện tuân thủ bảng màu Slate đơn sắc có độ tương phản cao, chuyển đổi Light/Dark mode với độ trễ 0ms.
• NFR-06 [Tính toàn vẹn dữ liệu]: Bật cơ chế PRAGMA foreign_keys = ON trên toàn bộ kết nối cơ sở dữ liệu để đảm bảo các ràng buộc khóa ngoại và xóa xếp tầng (CASCADE) hoạt động chính xác.

1.6. Xác định bài toán ứng dụng AI và phạm vi tích hợp
Trong khuôn khổ đề tài, Koshi tập trung ứng dụng các mô hình ngôn ngữ lớn (LLMs) vào 3 bài toán xử lý ngôn ngữ tự nhiên then chốt:
1. Bài toán Tóm tắt văn bản có cấu trúc (Structured Summarization): Tổng hợp dữ liệu trạng thái công việc trong sprint thành báo cáo súc tích 3 phần (Tổng quan, Điểm nghẽn tiến độ, Ưu tiên tiếp theo).
2. Bài toán Trích xuất thông tin thực thể (Information Extraction): Phân tích văn bản thô từ biên bản cuộc họp để trích xuất danh sách công việc, người chịu trách nhiệm, độ ưu tiên và thời hạn bàn giao.
3. Bài toán Tối ưu hóa phân bổ nguồn lực (Capacity Optimization): Đánh giá ma trận kỹ năng và số điểm độ phức tạp công việc đang thực hiện (WIP points) để đề xuất lập trình viên phù hợp nhất."""

doc.paragraphs[38].text = "CHƯƠNG 1. PHÂN TÍCH YÊU CẦU HỆ THỐNG"
doc.paragraphs[39].text = ch1_full_content
doc.paragraphs[40].text = ""
doc.paragraphs[41].text = ""
doc.paragraphs[42].text = ""

# Scope lock placeholders for Chapter 2, Chapter 3, and Conclusion
doc.paragraphs[43].text = "CHƯƠNG 2. THIẾT KẾ HỆ THỐNG (BÀI KIỂM TRA 2)"
doc.paragraphs[44].text = "[Nội dung thiết kế kiến trúc và CSDL chi tiết sẽ được hoàn thiện trong Bài kiểm tra 2 - KT2]"
doc.paragraphs[45].text = ""
doc.paragraphs[46].text = ""

doc.paragraphs[47].text = "CHƯƠNG 3. TÍCH HỢP AI VÀ ĐÁNH GIÁ (BÀI KIỂM TRA 3)"
doc.paragraphs[48].text = "[Nội dung kết quả thử nghiệm AI và triển khai thực tế sẽ được hoàn thiện trong Bài kiểm tra 3 - KT3]"
doc.paragraphs[49].text = ""
doc.paragraphs[50].text = ""
doc.paragraphs[51].text = ""
doc.paragraphs[52].text = ""
doc.paragraphs[53].text = ""

doc.paragraphs[54].text = "KẾT LUẬN"
doc.paragraphs[55].text = "[Kết luận toàn diện dự án sẽ được tổng kết sau khi hoàn thành các giai đoạn thử nghiệm KT2/KT3]"
doc.paragraphs[56].text = ""
doc.paragraphs[57].text = ""
doc.paragraphs[58].text = ""
doc.paragraphs[59].text = ""

print("[6/6] Updating References and Saving Document...")
doc.paragraphs[60].text = "TÀI LIỆU THAM KHẢO"
doc.paragraphs[61].text = "[1] ISO/IEC/IEEE 29148:2018, 'Systems and software engineering — Life cycle processes — Requirements engineering', IEEE Standards Association, 2018."
doc.paragraphs[62].text = "[2] Khoa Công nghệ Thông tin - Trường Đại học CNTT & Truyền thông Thái Nguyên (ICTU), 'Đề cương chi tiết và Quy chuẩn Đánh giá Dự án Học phần Ứng dụng Trí tuệ Nhân tạo', Thái Nguyên, 2026."
doc.paragraphs[63].text = "[3] Tiangolo, S., 'FastAPI Framework Documentation and Architecture Guidelines', Online: https://fastapi.tiangolo.com/, 2026."
doc.paragraphs[64].text = "[4] You, E. et al., 'Vue 3 Composition API & Pinia State Architecture Guide', Online: https://vuejs.org/, 2026."
doc.paragraphs[65].text = "[5] Google Cloud AI & DeepMind, 'Gemini API Technical Guidelines and Prompt Engineering Matrix', Online: https://ai.google.dev/, 2026."
doc.paragraphs[66].text = ""

doc.save(OUTPUT_PATH)
print(f"[✓] Document strictly mutated and saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    pass
