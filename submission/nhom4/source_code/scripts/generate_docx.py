import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_border(cell, **kwargs):
    """Set cell borders: top, bottom, left, right"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'<w:top w:val="{kwargs.get("top", "single")}" w:sz="{kwargs.get("top_sz", "4")}" w:space="0" w:color="{kwargs.get("top_color", "D3D3D3")}"/>\n'
        f'<w:left w:val="{kwargs.get("left", "single")}" w:sz="{kwargs.get("left_sz", "4")}" w:space="0" w:color="{kwargs.get("left_color", "D3D3D3")}"/>\n'
        f'<w:bottom w:val="{kwargs.get("bottom", "single")}" w:sz="{kwargs.get("bottom_sz", "4")}" w:space="0" w:color="{kwargs.get("bottom_color", "D3D3D3")}"/>\n'
        f'<w:right w:val="{kwargs.get("right", "single")}" w:sz="{kwargs.get("right_sz", "4")}" w:space="0" w:color="{kwargs.get("right_color", "D3D3D3")}"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def set_cell_shading(cell, color_hex):
    """Set background color of a cell"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def build_report():
    doc = Document()

    # --- PAGE SETUP (Standard A4, Margins: Top/Bottom/Right: 2cm, Left: 3cm) ---
    for section in doc.sections:
        section.top_margin = Inches(0.79)     # ~2.0 cm
        section.bottom_margin = Inches(0.79)  # ~2.0 cm
        section.left_margin = Inches(1.18)    # ~3.0 cm
        section.right_margin = Inches(0.79)   # ~2.0 cm

    # --- DEFAULT STYLES ---
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(13)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Slate-800
    normal_style.paragraph_format.line_spacing = 1.3
    normal_style.paragraph_format.space_after = Pt(6)

    # =========================================================================
    # 1. TRANG BÌA (COVER PAGE - MODEL THEO TEMPLATE ICTU)
    # =========================================================================
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_inst.add_run("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG\n")
    r1.bold = True
    r1.font.size = Pt(13)
    r2 = p_inst.add_run("KHOA CÔNG NGHỆ THÔNG TIN\n")
    r2.bold = True
    r2.font.size = Pt(13)
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(36)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_rep = p_title.add_run("BÁO CÁO DỰ ÁN\n")
    r_rep.bold = True
    r_rep.font.size = Pt(18)
    r_rep.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    r_sub = p_title.add_run("HỌC PHẦN ỨNG DỤNG TRÍ TUỆ NHÂN TẠO\n\n")
    r_sub.bold = True
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    p_topic = doc.add_paragraph()
    p_topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tlabel = p_topic.add_run("Đề tài:\n")
    r_tlabel.font.size = Pt(13)
    r_tname = p_topic.add_run("HỆ THỐNG QUẢN LÝ DỰ ÁN VÀ TIẾN ĐỘ CÔNG VIỆC KOSHI\nCÓ TÍCH HỢP TRÍ TUỆ NHÂN TẠO (AI)\n\n")
    r_tname.bold = True
    r_tname.font.size = Pt(15)

    r_grp = p_topic.add_run("Tên nhóm: NHÓM 04\n\n")
    r_grp.bold = True
    r_grp.font.size = Pt(14)

    # Table of Authors
    table_auth = doc.add_table(rows=5, cols=2)
    table_auth.alignment = WD_TABLE_ALIGNMENT.CENTER
    auth_data = [
        ("Nhóm sinh viên thực hiện:", ""),
        ("1. Phạm Minh Tú (Trưởng nhóm)", "MSSV: 23ICTU..."),
        ("2. Phạm Văn Huynh", "MSSV: 23ICTU..."),
        ("3. Đàm Đức Đôn", "MSSV: 23ICTU..."),
        ("Giảng viên hướng dẫn:", "ThS. Giảng Viên Phụ Trách")
    ]
    for idx, (col1, col2) in enumerate(auth_data):
        row = table_auth.rows[idx]
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.line_spacing = 1.2
        r = p0.add_run(col1)
        if idx == 0 or idx == 4:
            r.bold = True
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.line_spacing = 1.2
        p1.add_run(col2)
        row.cells[0].width = Inches(3.2)
        row.cells[1].width = Inches(2.3)

    p_bot = doc.add_paragraph()
    p_bot.paragraph_format.space_before = Pt(72)
    p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bot = p_bot.add_run("THÁI NGUYÊN, NĂM 2026")
    r_bot.bold = True
    r_bot.font.size = Pt(13)

    doc.add_page_break()

    # =========================================================================
    # 2. MỤC LỤC & PHÂN CÔNG NHIỆM VỤ
    # =========================================================================
    p_toc_head = doc.add_paragraph()
    p_toc_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_toc_head.add_run("MỤC LỤC")
    r.bold = True
    r.font.size = Pt(16)

    toc_items = [
        ("PHÂN CÔNG NHIỆM VỤ", "ii"),
        ("MỞ ĐẦU", "iv"),
        ("CHƯƠNG 1. PHÂN TÍCH YÊU CẦU HỆ THỐNG", "1"),
        ("  1.1. Bối cảnh bài toán và lý do phát triển", "1"),
        ("  1.2. Mô hình Actor và Phân hệ Use Case", "2"),
        ("  1.3. Yêu cầu chức năng (Functional Requirements)", "3"),
        ("  1.4. Yêu cầu phi chức năng (Non-Functional Requirements)", "4"),
        ("CHƯƠNG 2. THIẾT KẾ HỆ THỐNG VÀ CƠ SỞ DỮ LIỆU", "5"),
        ("  2.1. Kiến trúc tổng thể hệ thống", "5"),
        ("  2.2. Thiết kế Cơ sở dữ liệu quan hệ (ERD & Data Dictionary)", "6"),
        ("  2.3. Thiết kế trạng thái và toán học chuyển dịch trạng thái", "8"),
        ("CHƯƠNG 3. TÍCH HỢP AI, THIẾT KẾ PROMPT VÀ ĐÁNH GIÁ", "9"),
        ("  3.1. Xác định vị trí tích hợp AI trong hệ thống", "9"),
        ("  3.2. Thiết kế Prompt và Luồng gọi AI đa tầng (Cascade Architecture)", "10"),
        ("  3.3. Ma trận tối ưu hóa Prompt qua 3 vòng thử nghiệm", "12"),
        ("  3.4. Minh chứng ứng dụng AI trong quy trình SDLC (KT1)", "13"),
        ("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "14"),
        ("TÀI LIỆU THAM KHẢO", "15")
    ]
    for title, pg in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(3)
        r_t = p.add_run(title)
        if title.startswith("CHƯƠNG") or title in ["PHÂN CÔNG NHIỆM VỤ", "MỞ ĐẦU", "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "TÀI LIỆU THAM KHẢO"]:
            r_t.bold = True
        r_dots = p.add_run(" " + "." * (65 - len(title)) + " ")
        r_dots.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        r_p = p.add_run(pg)
        r_p.bold = True

    doc.add_page_break()

    # --- PHÂN CÔNG NHIỆM VỤ ---
    p_pc_head = doc.add_paragraph()
    p_pc_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_pc_head.add_run("PHÂN CÔNG NHIỆM VỤ")
    r.bold = True
    r.font.size = Pt(16)

    p_t1_title = doc.add_paragraph()
    r = p_t1_title.add_run("Bảng 1: Phân công nhiệm vụ theo tiến độ thực hiện (KT1)")
    r.bold = True

    t_progress = doc.add_table(rows=6, cols=4)
    t_progress.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_p = ["STT", "Nội dung công việc", "Người thực hiện", "Kết quả hoàn thành"]
    for i, h in enumerate(headers_p):
        cell = t_progress.rows[0].cells[i]
        set_cell_shading(cell, "F1F5F9")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True

    progress_rows = [
        ("1", "Phân tích bài toán, lập tài liệu URD, SRS, thiết kế Use Case & Actor", "Phạm Minh Tú", "100%"),
        ("2", "Thiết kế CSDL quan hệ (ERD, DDL schema.sql, init_db.py seeder)", "Phạm Văn Huynh", "100%"),
        ("3", "Thiết kế kiến trúc hệ thống 3 lớp & Luồng xử lý đồ thị DAG (Kahn)", "Phạm Minh Tú", "100%"),
        ("4", "Thiết kế System Prompt AI (Summary, Minutes, Workload) & Fallback", "Đàm Đức Đôn", "100%"),
        ("5", "Tổng hợp báo cáo kỹ thuật KT1, xây dựng DOCX & đóng gói submission", "Cả nhóm", "100%")
    ]
    for r_idx, row_data in enumerate(progress_rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = t_progress.rows[r_idx].cells[c_idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            if c_idx in [0, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    p_t2_title = doc.add_paragraph()
    r = p_t2_title.add_run("Bảng 2: Phân công nhiệm vụ theo thành viên thực hiện")
    r.bold = True

    t_members = doc.add_table(rows=4, cols=4)
    t_members.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_m = ["STT", "Thành viên", "Nhiệm vụ đảm nhiệm", "Ký tên"]
    for i, h in enumerate(headers_m):
        cell = t_members.rows[0].cells[i]
        set_cell_shading(cell, "F1F5F9")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True

    member_rows = [
        ("1", "Phạm Minh Tú\n(Trưởng nhóm)", "Chịu trách nhiệm kiến trúc tổng thể, mô hình toán học DAG/CPM, URD/SRS ISO standard, điều phối kỹ thuật.", ""),
        ("2", "Phạm Văn Huynh", "Phụ trách thiết kế mô hình CSDL, bảng từ điển dữ liệu, kiểm thử ràng buộc SQL và viết kịch bản test backend.", ""),
        ("3", "Đàm Đức Đôn", "Phụ trách xây dựng và tối ưu ma trận Prompt AI 3 vòng, thiết kế kịch bản test AI, tài liệu báo cáo.", "")
    ]
    for r_idx, row_data in enumerate(member_rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = t_members.rows[r_idx].cells[c_idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            if c_idx in [0, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_page_break()

    # =========================================================================
    # 3. MỞ ĐẦU
    # =========================================================================
    p_intro_head = doc.add_paragraph()
    p_intro_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_intro_head.add_run("MỞ ĐẦU")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph(
        "Trong bối cảnh các quy trình phát triển phần mềm hiện đại đòi hỏi tốc độ triển khai cực cao, các công cụ quản lý dự án truyền thống như Jira, Trello hay Asana đang bộc lộ những rào cản lớn về hiệu năng: thời gian phản hồi giao diện chậm, phụ thuộc hoàn toàn vào kết nối mạng, các biểu mẫu nhập liệu đa tầng gây gián đoạn luồng làm việc của lập trình viên, và thiếu sự hỗ trợ tự động hóa thông minh trong việc tổng hợp tiến độ cũng như phân bổ công việc."
    )
    doc.add_paragraph(
        "Hệ thống quản lý dự án Koshi (輿) được thiết kế và xây dựng nhằm giải quyết triệt để các hạn chế trên. Với triết lý 'Local-First' và điều hướng tối ưu hóa bàn phím (Vim-centric ergonomics), Koshi mang lại trải nghiệm tương tác với độ trễ dưới 16ms. Đồng thời, việc tích hợp Trí tuệ nhân tạo (AI) tạo sinh đóng vai trò như một trợ lý quản lý dự án tự động, hỗ trợ tóm tắt báo cáo tuần, bóc tách biên bản cuộc họp thành công việc cụ thể và đề xuất phân bổ nhân sự cân bằng tải."
    )
    doc.add_paragraph(
        "Báo cáo Bài kiểm tra 1 (KT1) trình bày toàn bộ kết quả phân tích yêu cầu, thiết kế kiến trúc hệ thống, mô hình hóa cơ sở dữ liệu và kế hoạch tích hợp AI của Nhóm 04. Chúng em xin chân thành cảm ơn giảng viên bộ môn đã tận tình hướng dẫn và định hướng để nhóm hoàn thiện tốt đề tài này."
    )

    doc.add_page_break()

    # =========================================================================
    # 4. CHƯƠNG 1: PHÂN TÍCH YÊU CẦU HỆ THỐNG
    # =========================================================================
    h1 = doc.add_paragraph()
    r = h1.add_run("CHƯƠNG 1. PHÂN TÍCH YÊU CẦU HỆ THỐNG")
    r.bold = True
    r.font.size = Pt(15)

    doc.add_heading("1.1. Bối cảnh bài toán và lý do phát triển", level=2)
    doc.add_paragraph(
        "Koshi được định vị là hệ thống quản lý công việc và tiến độ sprint nội bộ cho các nhóm kỹ sư phần mềm chuyên sâu. Bối cảnh nghiệp vụ tập trung vào 3 nhóm đối tượng chính (User Personas):"
    )
    p = doc.add_paragraph()
    p.add_run("• System Architect / Senior Engineer: ").bold = True
    p.add_run("Đòi hỏi thao tác bàn phím 100% không dùng chuột (h/j/k/l, Space, i, n, Esc), chuyển đổi nhanh giữa dạng Bảng (Table) và Kanban 2D, phát hiện sớm các điểm nghẽn tiến độ (Critical Path) trên đồ thị phụ thuộc.")
    
    p = doc.add_paragraph()
    p.add_run("• Project Manager (PM) / Tech Lead: ").bold = True
    p.add_run("Cần tự động hóa khâu lập báo cáo tiến độ tuần, tự động trích xuất đầu việc từ biên bản họp văn bản thô và nhận gợi ý phân công công việc dựa trên số điểm độ phức tạp (WIP story points).")
    
    p = doc.add_paragraph()
    p.add_run("• Field / Mobile Developer: ").bold = True
    p.add_run("Cần khả năng làm việc ngoại tuyến (Offline) khi mất kết nối mạng và đồng bộ tự động khi có mạng trở lại thông qua cơ chế lưu trữ IndexedDB cục bộ.")

    doc.add_heading("1.2. Mô hình Actor và Phân hệ Use Case", level=2)
    doc.add_paragraph("Hệ thống phân cấp 3 nhóm Actor với các quyền hạn nghiệp vụ tương ứng:")
    doc.add_paragraph("1. Guest: Đăng nhập hệ thống (Email/Password hoặc Google OAuth2), xem demo dữ liệu công khai.")
    doc.add_paragraph("2. Team Member (Lập trình viên): Duyệt bảng công việc, chuyển trạng thái vòng tròn (Space), chỉnh sửa chi tiết (i/Esc), tạo công việc mới (n), phân rã công việc bằng AI, bóc tách Git Diff.")
    doc.add_paragraph("3. Project Manager (Quản lý dự án): Toàn quyền quản trị thành viên, cập nhật kỹ năng (skills), sinh Báo cáo tổng kết tuần (AI Weekly Summary), bóc tách biên bản họp (AI Meeting Minutes), cân bằng tải nhân sự (AI Smart Assignment) và trực quan hóa đồ thị phụ thuộc (DAG).")

    doc.add_heading("1.3. Yêu cầu chức năng cốt lõi (Functional Requirements)", level=2)
    fr_items = [
        ("FR-01 [Dual-Mode Views]", "Chuyển đổi giao diện tức thì giữa Table View mật độ cao và Kanban Board 2D qua phím tắt 'b'."),
        ("FR-02 [Vim Ergonomics]", "Hỗ trợ điều hướng bàn phím đầy đủ: j/k duyệt dòng, h/j/k/l duyệt lưới Kanban, Space đổi trạng thái tuần hoàn."),
        ("FR-03 [Global Escape]", "Phím Escape ưu tiên bắt sự kiện ở capture-phase để đóng modal, hủy chế độ sửa và xóa focus input."),
        ("FR-04 [Local-First Persistence]", "Mọi thao tác thay đổi dữ liệu đều được ghi trực tiếp vào IndexedDB (idb-keyval) trước khi gọi API backend."),
        ("FR-05 [Topological DAG & CPM]", "Phân tích chuỗi phụ thuộc bằng giải thuật Kahn và tô màu nhận diện đường găng (Critical Path)."),
        ("FR-06 [AI PM Workflows]", "Tích hợp 3 chức năng AI: Tóm tắt tiến độ tuần, Bóc tách biên bản họp và Gợi ý phân bổ công việc.")
    ]
    for fid, fdesc in fr_items:
        p = doc.add_paragraph()
        r = p.add_run(f"• {fid}: ")
        r.bold = True
        p.add_run(fdesc)

    doc.add_page_break()

    # =========================================================================
    # 5. CHƯƠNG 2: THIẾT KẾ HỆ THỐNG VÀ CƠ SỞ DỮ LIỆU
    # =========================================================================
    h2 = doc.add_paragraph()
    r = h2.add_run("CHƯƠNG 2. THIẾT KẾ HỆ THỐNG VÀ CƠ SỞ DỮ LIỆU")
    r.bold = True
    r.font.size = Pt(15)

    doc.add_heading("2.1. Kiến trúc tổng thể hệ thống", level=2)
    doc.add_paragraph(
        "Koshi được thiết kế theo mô hình kiến trúc phân lớp hiện đại (3-Tier Decoupled Architecture), tối ưu cho tính sẵn sàng cao và bảo mật nghiêm ngặt:"
    )
    doc.add_paragraph("• Frontend Client: Xây dựng trên nền tảng Vue 3.5 (Composition API, `<script setup lang='ts'>`), quản lý trạng thái tập trung với Pinia 2.3, Vite 6 và hệ thống thiết kế giao diện tối giản Slate bằng Tailwind CSS v4.")
    doc.add_paragraph("• Edge Proxy & Web Server: Caddy Reverse Proxy tự động cấp phát chứng chỉ SSL/TLS, định tuyến lưu lượng vào Nginx Alpine đóng gói ứng dụng Single Page Application (SPA).")
    doc.add_paragraph("• Backend API Core: Sử dụng FastAPI (Python 3.11), SQLAlchemy 2.0 ORM, xác thực phân quyền qua OAuth2 Password Bearer & JSON Web Token (JWT HS256).")
    doc.add_paragraph("• Database Engine: Cơ sở dữ liệu quan hệ SQLite (Volume mount `/app/data/koshi.db`), đảm bảo tính gọn nhẹ, di động và hiệu năng truy vấn tức thì.")

    doc.add_heading("2.2. Thiết kế Cơ sở dữ liệu quan hệ (ERD & Data Dictionary)", level=2)
    doc.add_paragraph("CSDL bao gồm 6 bảng thực thể quan hệ chặt chẽ:")
    
    t_erd = doc.add_table(rows=7, cols=4)
    t_erd.alignment = WD_TABLE_ALIGNMENT.CENTER
    erd_headers = ["Tên bảng", "Khóa chính (PK)", "Khóa ngoại (FK)", "Mô tả nghiệp vụ"]
    for i, h in enumerate(erd_headers):
        cell = t_erd.rows[0].cells[i]
        set_cell_shading(cell, "F1F5F9")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True

    erd_rows = [
        ("users", "id (INTEGER)", "Không", "Lưu thông tin tài khoản, mật khẩu băm Bcrypt, vai trò (PM/MEMBER), kỹ năng (skills)."),
        ("projects", "id (INTEGER)", "owner_id -> users(id)", "Lưu thông tin dự án phần mềm và chủ sở hữu dự án."),
        ("sprints", "id (INTEGER)", "project_id -> projects(id)", "Quản lý chu kỳ sprint, mục tiêu và trạng thái kích hoạt."),
        ("tasks", "id (VARCHAR)", "project_id, sprint_id, assignee_id", "Lưu chi tiết công việc, trạng thái, độ ưu tiên, điểm phức tạp (1,2,3,5)."),
        ("task_dependencies", "id (INTEGER)", "task_id, depends_on_task_id", "Lưu quan hệ phụ thuộc giữa các đầu việc phục vụ giải thuật DAG."),
        ("comments", "id (INTEGER)", "task_id, author_id", "Lưu lịch sử trao đổi và bình luận kỹ thuật của thành viên.")
    ]
    for r_idx, row_data in enumerate(erd_rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = t_erd.rows[r_idx].cells[c_idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            if c_idx in [0, 1]:
                p.add_run(val).bold = True
            else:
                p.add_run(val)

    doc.add_page_break()

    # =========================================================================
    # 6. CHƯƠNG 3: TÍCH HỢP AI, THIẾT KẾ PROMPT VÀ ĐÁNH GIÁ
    # =========================================================================
    h3 = doc.add_paragraph()
    r = h3.add_run("CHƯƠNG 3. TÍCH HỢP AI, THIẾT KẾ PROMPT VÀ ĐÁNH GIÁ")
    r.bold = True
    r.font.size = Pt(15)

    doc.add_heading("3.1. Kiến trúc luồng gọi AI đa tầng (AI Multi-Tier Cascade)", level=2)
    doc.add_paragraph(
        "Nhằm đảm bảo hệ thống không bao giờ bị gián đoạn hoạt động khi API bên ngoài gặp sự cố hoặc cạn kiệt hạn ngạch (Rate Limit), Koshi thiết kế cơ chế dự phòng 3 tầng (Cascade Fallback):"
    )
    doc.add_paragraph("• Tầng 1 (Primary): Gọi trực tiếp API mô hình ngôn ngữ lớn thương mại (OpenAI GPT-4o / Google Gemini 1.5 Pro) để đạt chất lượng suy luận cao nhất.")
    doc.add_paragraph("• Tầng 2 (Secondary Local): Tự động chuyển tiếp sang endpoint máy chủ cục bộ Ollama (chạy model qwen2.5-coder hoặc llama3.2) khi mất kết nối Internet.")
    doc.add_paragraph("• Tầng 3 (Deterministic Heuristic Fallback): Phân tích cú pháp AST và áp dụng biểu thức chính quy (Regex Rule Engine) trực tiếp trên máy chủ để luôn trả về kết quả JSON hợp lệ 100% mà không bị lỗi crash.")

    doc.add_heading("3.2. Ma trận tối ưu hóa Prompt qua 3 vòng thử nghiệm", level=2)
    
    t_prompt = doc.add_table(rows=4, cols=4)
    t_prompt.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_headers = ["Chức năng AI", "Vòng 1 (Zero-Shot)", "Vòng 2 (Ràng buộc cấu trúc)", "Vòng 3 (Strict JSON & Fallback)"]
    for i, h in enumerate(p_headers):
        cell = t_prompt.rows[0].cells[i]
        set_cell_shading(cell, "F1F5F9")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True

    p_rows = [
        ("Weekly Summary", "Văn bản thô dài dòng, hallucinate đầu việc không có thật.", "Chia 3 mục cố định (Overview, Blockers, Priorities), còn sót markdown.", "Chuẩn hóa định dạng Schema Pydantic, tự động nhận diện thẻ [Critical Path]."),
        ("Meeting Minutes", "Liệt kê gạch đầu dòng, bỏ sót người chịu trách nhiệm và deadline.", "Bắt buộc trả JSON schema, thi thoảng sinh lỗi cú pháp khi transcript dài.", "Tích hợp bộ parse JSON tự sửa lỗi + Fallback Regex trích xuất Speaker."),
        ("Smart Workload", "Gợi ý cảm tính, không tính toán dung lượng tải hiện tại.", "Bổ sung điểm Story Point, chưa giải quyết được trường hợp hòa điểm.", "Áp dụng hàm tối ưu hóa chi phí J(u) kết hợp chặn trần quá tải 8 điểm WIP.")
    ]
    for r_idx, row_data in enumerate(p_rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = t_prompt.rows[r_idx].cells[c_idx]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.add_run(val)

    doc.add_page_break()

    # =========================================================================
    # 7. KẾT LUẬN & TÀI LIỆU THAM KHẢO
    # =========================================================================
    p_conc = doc.add_paragraph()
    p_conc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_conc.add_run("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph(
        "Giai đoạn Bài kiểm tra 1 (KT1) đã hoàn thành xuất sắc toàn bộ các mục tiêu đặt ra: phân tích thấu đáo bài toán quản lý dự án, thiết kế kiến trúc 3 lớp vững chắc, hoàn thiện mô hình cơ sở dữ liệu quan hệ với các ràng buộc toàn vẹn dữ liệu, và xây dựng thành công cơ chế tích hợp AI đa tầng linh hoạt."
    )
    doc.add_paragraph(
        "Kế hoạch triển khai cho các giai đoạn tiếp theo:"
    )
    doc.add_paragraph("• Bài kiểm tra 2 (KT2): Hoàn thiện toàn bộ các API CRUD nghiệp vụ, phân quyền JWT RBAC, tích hợp giao diện phím tắt 2D Kanban và bảng thống kê tiến độ.")
    doc.add_paragraph("• Bài kiểm tra 3 (KT3): Tối ưu hóa sâu luồng gọi AI đa tầng, kiểm thử tự động toàn diện với Pytest và triển khai container hóa bằng Docker Compose trên hạ tầng máy chủ thực tế.")

    doc.add_paragraph().paragraph_format.space_before = Pt(24)

    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_ref.add_run("TÀI LIỆU THAM KHẢO")
    r.bold = True
    r.font.size = Pt(16)

    references = [
        "[1] ISO/IEC/IEEE 29148:2018, 'Systems and software engineering — Life cycle processes — Requirements engineering', IEEE Standards Association, 2018.",
        "[2] Tiêu chuẩn Đánh giá Dự án Học phần Ứng dụng Trí tuệ Nhân tạo, Khoa Công nghệ Thông tin - Trường Đại học CNTT & Truyền thông Thái Nguyên (ICTU), 2026.",
        "[3] FastAPI Framework Documentation, Online: [https://fastapi.tiangolo.com/](https://fastapi.tiangolo.com/), Truy cập tháng 08/2026.",
        "[4] Vue 3 Composition API & Pinia State Engine, Online: [https://vuejs.org/](https://vuejs.org/), Truy cập tháng 08/2026.",
        "[5] Google Gemini API & Prompt Engineering Guidelines, Online: [https://ai.google.dev/](https://ai.google.dev/), Truy cập tháng 08/2026."
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        p.add_run(ref)

    repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    output_path = os.path.join(repo_root, "nhom4.docx")
    doc.save(output_path)
    print(f"Report successfully compiled to: {output_path}")

if __name__ == "__main__":
    build_report()
